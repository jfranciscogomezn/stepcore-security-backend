"""Generate comprehensive StepCore UAT test plan Word document."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = (
    Path(__file__).resolve().parent.parent
    / "docs"
    / "qa"
    / "Plan-Pruebas-Usuario-StepCore.docx"
)


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = "D9E2F3",
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    doc.add_paragraph()


def add_test_case(
    doc: Document,
    case_id: str,
    title: str,
    priority: str,
    precondition: str,
    steps: list[str],
    expected: str,
    module: str = "",
) -> None:
    doc.add_heading(f"{case_id} — {title}", level=3)
    if module:
        p = doc.add_paragraph()
        p.add_run("Módulo: ").bold = True
        p.add_run(module)
    p = doc.add_paragraph()
    p.add_run("Prioridad: ").bold = True
    p.add_run(priority)
    p = doc.add_paragraph()
    p.add_run("Precondición: ").bold = True
    p.add_run(precondition)
    doc.add_paragraph("Pasos:")
    for step in steps:
        doc.add_paragraph(step, style="List Number")
    p = doc.add_paragraph()
    p.add_run("Resultado esperado: ").bold = True
    p.add_run(expected)
    doc.add_paragraph()


def add_module_intro(doc: Document, title: str, description: str, routes: list[str]) -> None:
    doc.add_heading(title, level=2)
    doc.add_paragraph(description)
    if routes:
        doc.add_paragraph("Rutas / pantallas principales:", style="List Bullet")
        for route in routes:
            doc.add_paragraph(route, style="List Bullet")


def add_cases(doc: Document, cases: list[dict], module: str) -> None:
    for case in cases:
        add_test_case(doc, module=module, **case)


# ---------------------------------------------------------------------------
# Test case definitions by module
# ---------------------------------------------------------------------------

AUTH_CASES = [
    {
        "case_id": "AUTH-001",
        "title": "Login exitoso con tenant slug",
        "priority": "Alta",
        "precondition": "Usuario maria.lopez@acme.qa activo; tenant acme.",
        "steps": [
            "Abrir /login.",
            "Ingresar slug acme, correo y contraseña Admin@2026!.",
            "Pulsar Iniciar sesión.",
        ],
        "expected": "Redirección al dashboard. Sidebar muestra menú según rol EMPLOYEE. Token JWT almacenado.",
    },
    {
        "case_id": "AUTH-002",
        "title": "Login fallido — contraseña incorrecta",
        "priority": "Alta",
        "precondition": "Cuenta válida.",
        "steps": [
            "Intentar login con contraseña errónea tres veces.",
        ],
        "expected": "Mensaje de error claro; no se expone si el correo existe. No hay sesión activa.",
    },
    {
        "case_id": "AUTH-003",
        "title": "Login fallido — tenant incorrecto",
        "priority": "Alta",
        "precondition": "Usuario existe solo en tenant acme.",
        "steps": [
            "Login con slug globex y credenciales de usuario acme.",
        ],
        "expected": "Autenticación rechazada (401). Usuario no accede a datos de otro tenant.",
    },
    {
        "case_id": "AUTH-004",
        "title": "Cierre de sesión",
        "priority": "Media",
        "precondition": "Sesión activa.",
        "steps": [
            "Usar acción Cerrar sesión en sidebar.",
            "Intentar navegar a /dashboard sin login.",
        ],
        "expected": "Redirección a /login. Rutas protegidas inaccesibles.",
    },
    {
        "case_id": "AUTH-005",
        "title": "Rutas protegidas por permiso",
        "priority": "Alta",
        "precondition": "Login como empleado (sin TIME_RECORDS_ADMIN).",
        "steps": [
            "Intentar acceder manualmente a /admin/time.",
        ],
        "expected": "Acceso denegado o redirección; no se muestra UI admin de tiempo.",
    },
    {
        "case_id": "AUTH-006",
        "title": "Cambio de contraseña obligatorio",
        "priority": "Media",
        "precondition": "Usuario con mustChangePassword=true (p. ej. admin recién aprovisionado).",
        "steps": [
            "Iniciar sesión.",
            "Completar formulario de cambio de contraseña en /my/profile.",
        ],
        "expected": "Contraseña actualizada; usuario puede continuar; política de complejidad validada en UI.",
    },
]

TEN_CASES = [
    {
        "case_id": "TEN-001",
        "title": "Listar inquilinos (platform admin)",
        "priority": "Alta",
        "precondition": "Login platform@stepcore.com / tenant platform / Platform@2026!.",
        "steps": [
            "Ir a /platform/tenants.",
            "Verificar listado incluye acme, globex, legacy.",
        ],
        "expected": "Tabla con nombre, slug, plan, estado y acciones. Sin datos de empleados de tenant mezclados.",
    },
    {
        "case_id": "TEN-002",
        "title": "Aprovisionar nuevo inquilino",
        "priority": "Alta",
        "precondition": "Platform admin autenticado.",
        "steps": [
            "Crear inquilino qa-test-XX con admin inicial.",
            "Anotar contraseña temporal mostrada una sola vez.",
            "Login con nuevo admin en slug creado.",
        ],
        "expected": "Tenant ACTIVE; rol ADMIN con menú completo; admin puede operar su tenant aislado.",
    },
    {
        "case_id": "TEN-003",
        "title": "Suspender inquilino",
        "priority": "Media",
        "precondition": "Tenant de prueba activo.",
        "steps": [
            "Suspender tenant desde plataforma.",
            "Intentar login como admin de ese tenant.",
        ],
        "expected": "Login rechazado. Tenant SUSPENDED en listado.",
    },
    {
        "case_id": "TEN-004",
        "title": "Admin tenant no accede a plataforma",
        "priority": "Alta",
        "precondition": "Login admin@acme.com.",
        "steps": [
            "Navegar a /platform/tenants.",
        ],
        "expected": "403 o ruta no accesible; menú de plataforma no visible.",
    },
    {
        "case_id": "TEN-005",
        "title": "Aislamiento de datos entre tenants",
        "priority": "Alta",
        "precondition": "Datos QA en acme y globex.",
        "steps": [
            "Admin acme: listar empleados y usuarios.",
            "Admin globex: repetir.",
        ],
        "expected": "Cada tenant ve únicamente sus registros (@acme.qa vs @globex.qa).",
    },
]

ACC_CASES = [
    {
        "case_id": "ACC-001",
        "title": "Hub de control de acceso",
        "priority": "Media",
        "precondition": "Admin tenant autenticado.",
        "steps": [
            "Ir a /admin/access.",
            "Abrir enlaces a catálogo, roles y usuarios.",
        ],
        "expected": "Tres tarjetas navegables; textos coherentes con permisos ADMIN.",
    },
    {
        "case_id": "ACC-002",
        "title": "CRUD de rol y permisos de menú",
        "priority": "Alta",
        "precondition": "Admin tenant.",
        "steps": [
            "Crear rol SUPERVISOR_QA.",
            "En pestaña Permisos, asignar MY_TIME y REPORTS.",
            "Guardar permisos.",
        ],
        "expected": "Rol persistido; solo ítems hoja seleccionados otorgan acceso.",
    },
    {
        "case_id": "ACC-003",
        "title": "CRUD de usuario y asignación de rol",
        "priority": "Alta",
        "precondition": "Rol SUPERVISOR_QA existente.",
        "steps": [
            "Crear usuario qa.supervisor@acme.qa con rol SUPERVISOR_QA.",
            "Login con ese usuario.",
        ],
        "expected": "Menú lateral refleja solo pantallas asignadas al rol.",
    },
    {
        "case_id": "ACC-004",
        "title": "Desactivar usuario",
        "priority": "Media",
        "precondition": "Usuario de prueba activo.",
        "steps": [
            "Desactivar usuario desde detalle.",
            "Intentar login.",
        ],
        "expected": "Login rechazado; usuario aparece como inactivo en listado.",
    },
    {
        "case_id": "ACC-005",
        "title": "Restablecer contraseña (admin)",
        "priority": "Media",
        "precondition": "Admin y usuario objetivo.",
        "steps": [
            "Ejecutar Restablecer contraseña desde detalle de usuario.",
            "Login con nueva contraseña temporal/policy.",
        ],
        "expected": "Operación exitosa; auditoría registrada si aplica.",
    },
    {
        "case_id": "ACC-006",
        "title": "Catálogo de menú (platform admin)",
        "priority": "Baja",
        "precondition": "Platform admin.",
        "steps": [
            "Abrir /admin/access/menu.",
            "Ver árbol MODULE → GROUP → ITEM.",
        ],
        "expected": "Estructura jerárquica visible; edición según permisos de plataforma.",
    },
]

PAY_CASES = [
    {
        "case_id": "PAY-001",
        "title": "Consultar configuración de nómina del año",
        "priority": "Alta",
        "precondition": "Admin acme; payroll 2026 sembrado.",
        "steps": [
            "Ir a /admin/config.",
            "Seleccionar año 2026.",
        ],
        "expected": "Formulario cargado: salario mínimo, subsidio transporte, horas, factores de recargo, rangos horarios.",
    },
    {
        "case_id": "PAY-002",
        "title": "Crear configuración para año nuevo",
        "priority": "Media",
        "precondition": "Año sin config (p. ej. 2027).",
        "steps": [
            "Seleccionar 2027.",
            "Completar campos obligatorios con valores válidos.",
            "Guardar.",
        ],
        "expected": "Config creada (201); mensaje de éxito; recarga muestra valores guardados.",
    },
    {
        "case_id": "PAY-003",
        "title": "Actualizar factores de recargo",
        "priority": "Media",
        "precondition": "Config 2026 existente.",
        "steps": [
            "Modificar factor horas extra diurnas.",
            "Guardar y recargar página.",
        ],
        "expected": "Cambio persistido; reportes posteriores usan nuevos factores.",
    },
    {
        "case_id": "PAY-004",
        "title": "Agregar festivo",
        "priority": "Alta",
        "precondition": "Config año activo.",
        "steps": [
            "Agregar festivo 2026-12-08 con descripción.",
            "Verificar aparece en listado de festivos.",
        ],
        "expected": "Festivo guardado; clasificación de devengados trata esa fecha como festivo.",
    },
    {
        "case_id": "PAY-005",
        "title": "Validación de campos obligatorios",
        "priority": "Media",
        "precondition": "Formulario vacío o año nuevo.",
        "steps": [
            "Intentar guardar con salario mínimo en cero o campos vacíos.",
        ],
        "expected": "Validación en UI/API; no se persiste config inválida.",
    },
]

EMP_CASES = [
    {
        "case_id": "EMP-001",
        "title": "Listar empleados",
        "priority": "Alta",
        "precondition": "Admin acme; seed aplicado.",
        "steps": [
            "Ir a /admin/employees.",
        ],
        "expected": "Tabla con 4 empleados @acme.qa; columnas nombre, documento, correo, salario.",
    },
    {
        "case_id": "EMP-002",
        "title": "Crear empleado",
        "priority": "Alta",
        "precondition": "Admin tenant.",
        "steps": [
            "Nuevo empleado: datos personales, CC, salario > 0.",
            "Opcional: vincular userId de usuario existente.",
            "Guardar.",
        ],
        "expected": "Empleado creado; aparece en listado; documento y email únicos por tenant.",
    },
    {
        "case_id": "EMP-003",
        "title": "Editar empleado",
        "priority": "Media",
        "precondition": "Empleado existente.",
        "steps": [
            "Modificar teléfono y salario.",
            "Guardar.",
        ],
        "expected": "Cambios visibles en listado y detalle.",
    },
    {
        "case_id": "EMP-004",
        "title": "Duplicado documento / email rechazado",
        "priority": "Media",
        "precondition": "maria.lopez@acme.qa existente.",
        "steps": [
            "Intentar crear otro empleado con mismo email o idNumber.",
        ],
        "expected": "Error 409 o mensaje de duplicado; no se crea registro.",
    },
    {
        "case_id": "EMP-005",
        "title": "Empleado vinculado a usuario — clock-in",
        "priority": "Alta",
        "precondition": "Empleado con user_id y usuario EMPLOYEE.",
        "steps": [
            "Login maria.lopez@acme.qa.",
            "Registrar entrada en Mi tiempo (si no hay registro hoy).",
        ],
        "expected": "Clock-in asociado al perfil empleado correcto; no error de perfil no vinculado.",
    },
]

TIME_CASES = [
    {
        "case_id": "TIME-001",
        "title": "Registrar entrada (clock-in)",
        "priority": "Alta",
        "precondition": "Empleado sin registro OPEN hoy (usar usuario de prueba o día nuevo).",
        "steps": [
            "Mi tiempo → Registrar entrada.",
        ],
        "expected": "Registro OPEN creado; hora de entrada visible; botón salida habilitado.",
    },
    {
        "case_id": "TIME-002",
        "title": "Registrar salida (clock-out)",
        "priority": "Alta",
        "precondition": "Registro OPEN del día.",
        "steps": [
            "Registrar salida.",
        ],
        "expected": "Estado CLOSED; duración calculada; botones actualizados.",
    },
    {
        "case_id": "TIME-003",
        "title": "No permitir doble entrada el mismo día",
        "priority": "Alta",
        "precondition": "Ya existe registro para hoy.",
        "steps": [
            "Intentar segundo clock-in.",
        ],
        "expected": "Botón deshabilitado o error API; un solo registro por work_date.",
    },
    {
        "case_id": "TIME-004",
        "title": "Historial con devengados (capped)",
        "priority": "Alta",
        "precondition": "Maria Lopez; sin incompletos bloqueantes en rango.",
        "steps": [
            "Mi tiempo → rango Este mes.",
            "Revisar columnas Normal, H.E. diurna, Devengado.",
        ],
        "expected": "Filas CLOSED muestran minutos clasificados y montos; total del periodo visible.",
    },
    {
        "case_id": "TIME-005",
        "title": "Historial bloqueado por incompletos",
        "priority": "Alta",
        "precondition": "Ana Torres (tiene INCOMPLETE).",
        "steps": [
            "Mi tiempo → historial mayo 2026.",
        ],
        "expected": "Alerta: resumen de devengados oculto; columnas de earnings no mostradas hasta resolver incompletos.",
    },
    {
        "case_id": "TIME-006",
        "title": "Banner empleado — registros incompletos",
        "priority": "Alta",
        "precondition": "Ana Torres autenticada.",
        "steps": [
            "Dashboard o Mi tiempo.",
        ],
        "expected": "Banner warning con fechas incompletas y enlace a Mi tiempo.",
    },
    {
        "case_id": "TIME-007",
        "title": "Presets de periodo (mes, semana, 30 días)",
        "priority": "Baja",
        "precondition": "Mi tiempo abierto.",
        "steps": [
            "Pulsar Este mes, Esta semana, Últimos 30 días.",
        ],
        "expected": "Filtros from/to actualizados; tabla recarga acorde.",
    },
]

TADM_CASES = [
    {
        "case_id": "TADM-001",
        "title": "Listar registros por empleado y rango",
        "priority": "Alta",
        "precondition": "Admin acme.",
        "steps": [
            "/admin/time → seleccionar Maria Lopez → mayo 2026.",
        ],
        "expected": "Tabla con fechas, entradas/salidas, duración, estado, acciones.",
    },
    {
        "case_id": "TADM-002",
        "title": "Pestaña Incompletos",
        "priority": "Alta",
        "precondition": "Ana Torres con INCOMPLETE.",
        "steps": [
            "Pestaña Incompletos (con o sin filtro empleado).",
        ],
        "expected": "Registro 2026-05-19 listado; acción Resolver disponible.",
    },
    {
        "case_id": "TADM-003",
        "title": "Resolver registro incompleto",
        "priority": "Alta",
        "precondition": "Registro INCOMPLETE.",
        "steps": [
            "Resolver → ingresar salida manual y nota.",
            "Confirmar.",
        ],
        "expected": "Estado CLOSED; desaparece de incompletos; empleado puede ver devengados.",
    },
    {
        "case_id": "TADM-004",
        "title": "Corregir timestamps",
        "priority": "Alta",
        "precondition": "Registro CLOSED.",
        "steps": [
            "Corregir → ajustar entrada/salida + motivo.",
            "Guardar.",
        ],
        "expected": "Registro marcado corrected; auditoría generada; reportes recalculados.",
    },
    {
        "case_id": "TADM-005",
        "title": "Reabrir registro cerrado",
        "priority": "Media",
        "precondition": "Registro CLOSED reciente.",
        "steps": [
            "Reabrir registro.",
        ],
        "expected": "Estado OPEN; salida eliminada o según reglas de negocio; audit trail.",
    },
    {
        "case_id": "TADM-006",
        "title": "Crear registro manual (admin)",
        "priority": "Alta",
        "precondition": "Empleado seleccionado.",
        "steps": [
            "Crear registro → fecha, entrada, salida, motivo.",
        ],
        "expected": "Registro CLOSED creado; visible en listado y reportes.",
    },
]

RPT_CASES = [
    {
        "case_id": "RPT-001",
        "title": "Reporte propio (empleado) — vista con tope",
        "priority": "Alta",
        "precondition": "maria.lopez@acme.qa → /my/reports.",
        "steps": [
            "Periodo Este mes.",
        ],
        "expected": "Tabla con minutos clasificados capped y devengado; total del periodo.",
    },
    {
        "case_id": "RPT-002",
        "title": "Reporte admin por empleado",
        "priority": "Alta",
        "precondition": "Admin → /reports.",
        "steps": [
            "Seleccionar Maria Lopez; mes mayo 2026.",
        ],
        "expected": "Reporte generado; resumen nombre · periodo · total.",
    },
    {
        "case_id": "RPT-003",
        "title": "Vista sin tope (Report B)",
        "priority": "Media",
        "precondition": "Admin con empleado seleccionado.",
        "steps": [
            "Activar switch Uncapped / sin tope.",
        ],
        "expected": "Montos uncapped mayores o iguales a capped; columnas classified vs capped coherentes.",
    },
    {
        "case_id": "RPT-004",
        "title": "Bloqueo por registros incompletos (409)",
        "priority": "Alta",
        "precondition": "Ana Torres seleccionada.",
        "steps": [
            "Generar reporte que incluya fecha incompleta.",
        ],
        "expected": "Alerta UI; no tabla de resultados; fechas incompletas listadas.",
    },
    {
        "case_id": "RPT-005",
        "title": "Filtros de periodo — día, semana, rango",
        "priority": "Media",
        "precondition": "Admin o empleado en reportes.",
        "steps": [
            "Probar modos: mes, día único, semana, rango personalizado.",
            "Usar presets Hoy / Esta semana.",
        ],
        "expected": "Datos acotados al periodo; etiqueta de periodo correcta.",
    },
    {
        "case_id": "RPT-006",
        "title": "Resaltado horas extra / extendidas",
        "priority": "Media",
        "precondition": "Maria Lopez o Sofia Ramirez (días con overtime en seed).",
        "steps": [
            "Reporte mayo 2026.",
            "Identificar filas WARNING / ALERT.",
        ],
        "expected": "Badges Overtime / Extended hours; filas con clase highlight visual.",
    },
    {
        "case_id": "RPT-007",
        "title": "Exportar Excel",
        "priority": "Alta",
        "precondition": "Reporte cargado sin bloqueo.",
        "steps": [
            "Exportar Excel.",
            "Abrir archivo descargado.",
        ],
        "expected": "Archivo time-report.xlsx descargado; contiene registros y columna Notes si aplica.",
    },
    {
        "case_id": "RPT-008",
        "title": "Turno nocturno y festivo en clasificación",
        "priority": "Media",
        "precondition": "Carlos Ruiz (nocturno + festivo 25-may).",
        "steps": [
            "Reporte admin Carlos Ruiz mayo 2026.",
        ],
        "expected": "Minutos nocturnos/festivos > 0 en días correspondientes; devengado coherente con config.",
    },
]

AUD_CASES = [
    {
        "case_id": "AUD-001",
        "title": "Consultar historial de auditoría",
        "priority": "Alta",
        "precondition": "Acciones admin previas (corregir/crear/reabrir).",
        "steps": [
            "/admin/time/audit → rango mayo 2026.",
        ],
        "expected": "Entradas con cuándo, acción, usuario actor, empleado, antes/después.",
    },
    {
        "case_id": "AUD-002",
        "title": "Filtrar por empleado",
        "priority": "Media",
        "precondition": "Historial con múltiples empleados.",
        "steps": [
            "Filtro empleado = Maria Lopez → Actualizar.",
        ],
        "expected": "Solo entradas relacionadas con registros de ese empleado.",
    },
    {
        "case_id": "AUD-003",
        "title": "Filtrar por usuario actor",
        "priority": "Media",
        "precondition": "Varios admins han operado.",
        "steps": [
            "Filtro usuario actor = admin@acme.com.",
        ],
        "expected": "Solo cambios realizados por ese usuario.",
    },
    {
        "case_id": "AUD-004",
        "title": "Inmutabilidad — solo lectura",
        "priority": "Baja",
        "precondition": "Pantalla de auditoría.",
        "steps": [
            "Verificar ausencia de acciones editar/eliminar en filas.",
        ],
        "expected": "Auditoría es consultiva; no modifica registros desde UI.",
    },
]

NOT_CASES = [
    {
        "case_id": "NOT-001",
        "title": "Banner admin — registros incompletos",
        "priority": "Alta",
        "precondition": "Admin TIME_RECORDS_ADMIN; incompletos existentes (Ana/Paula).",
        "steps": [
            "Login admin@acme.com → Dashboard.",
        ],
        "expected": "Banner warning con título/mensaje de notificación y lista empleado — fecha; enlace a /admin/time.",
    },
    {
        "case_id": "NOT-002",
        "title": "Banner no visible para empleado",
        "priority": "Media",
        "precondition": "Login empleado sin permiso admin.",
        "steps": [
            "Dashboard.",
        ],
        "expected": "No aparece banner admin; sí puede aparecer banner propio de incompletos (TIME-006).",
    },
    {
        "case_id": "NOT-003",
        "title": "Actualización tras resolver incompleto",
        "priority": "Media",
        "precondition": "Banner admin visible.",
        "steps": [
            "Resolver último incompleto del tenant.",
            "Refrescar dashboard.",
        ],
        "expected": "Banner desaparece o reduce conteo tras resolución y refresh.",
    },
]

PROF_CASES = [
    {
        "case_id": "PROF-001",
        "title": "Ver información de cuenta",
        "priority": "Baja",
        "precondition": "Usuario autenticado.",
        "steps": [
            "/my/profile.",
        ],
        "expected": "Nombre, correo, rol visibles; sin editar campos protegidos.",
    },
    {
        "case_id": "PROF-002",
        "title": "Cambio de contraseña voluntario",
        "priority": "Media",
        "precondition": "Usuario con contraseña conocida.",
        "steps": [
            "Cambiar contraseña con política válida.",
        ],
        "expected": "Éxito; login con nueva contraseña funciona.",
    },
]

I18N_CASES = [
    {
        "case_id": "I18N-001",
        "title": "Locale predeterminado es-CO",
        "priority": "Alta",
        "precondition": "localStorage limpio.",
        "steps": ["Abrir /login.", "Observar labels."],
        "expected": "UI en español (Colombia).",
    },
    {
        "case_id": "I18N-002",
        "title": "Cambio de idioma en login",
        "priority": "Alta",
        "precondition": "Pantalla login.",
        "steps": ["Seleccionar English (US)."],
        "expected": "Textos en inglés sin recarga manual.",
    },
    {
        "case_id": "I18N-003",
        "title": "Persistencia de locale",
        "priority": "Alta",
        "precondition": "en-US seleccionado.",
        "steps": ["Login → F5 → navegar."],
        "expected": "Locale persiste en localStorage y UI.",
    },
    {
        "case_id": "I18N-004",
        "title": "Selector en sidebar",
        "priority": "Media",
        "precondition": "Sesión activa.",
        "steps": ["Cambiar idioma desde sidebar."],
        "expected": "Menú y pantalla actualizados.",
    },
    {
        "case_id": "I18N-005",
        "title": "Menú por código MY_TIME",
        "priority": "Alta",
        "precondition": "es-CO.",
        "steps": ["Ver sidebar empleado."],
        "expected": "Etiqueta traducida (Mi tiempo).",
    },
    {
        "case_id": "I18N-006",
        "title": "Formato moneda es-CO",
        "priority": "Alta",
        "precondition": "Reporte con montos.",
        "steps": ["Ver columnas Devengado."],
        "expected": "Formato Intl es-CO.",
    },
    {
        "case_id": "I18N-007",
        "title": "Formato moneda en-US",
        "priority": "Media",
        "precondition": "Locale en-US.",
        "steps": ["Abrir reportes."],
        "expected": "Encabezados y números en formato US.",
    },
    {
        "case_id": "I18N-008",
        "title": "Errores fallback frontend",
        "priority": "Media",
        "precondition": "API caída.",
        "steps": ["Provocar error de carga."],
        "expected": "Mensaje traducido en ApiErrorAlert.",
    },
]


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_heading("Plan de pruebas de usuario — StepCore", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Versión 2.0  |  Fecha: {date.today().isoformat()}  |  "
        "Alcance: módulos Payroll & Time implementados + i18n web"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    # 1. Introducción
    doc.add_heading("1. Introducción", level=1)
    doc.add_paragraph(
        "Este documento define las pruebas de aceptación de usuario (UAT) para la plataforma "
        "StepCore en su dominio Payroll & Time Management, incluyendo los módulos ya construidos "
        "y desplegados en entorno local: autenticación, multi-tenancy, control de acceso, "
        "configuración de nómina, empleados, registro de tiempo, cálculo y reportes de devengados, "
        "auditoría de tiempo, notificaciones administrativas e internacionalización del frontend web."
    )
    doc.add_paragraph(
        "Cada caso indica qué se prueba, cómo ejecutarlo paso a paso y el resultado esperado. "
        "Los datos de prueba sembrados (sección 6) permiten recorrer escenarios realistas sin "
        "configuración manual extensa."
    )
    doc.add_paragraph(
        "Tipo de prueba: manual funcional en navegador. Complemento recomendado: pruebas "
        "automatizadas de backend (IT) ya existentes en los repositorios security y business."
    )

    # 2. Módulos implementados
    doc.add_heading("2. Módulos implementados (alcance del documento)", level=1)
    add_table(
        doc,
        ["Módulo", "Estado", "Pantallas / capacidades", "IDs casos"],
        [
            ["Autenticación y sesión", "Completo", "Login, logout, rutas protegidas, perfil", "AUTH, PROF"],
            ["Multi-tenancy (SaaS)", "Completo", "Inquilinos, aprovisionamiento, aislamiento", "TEN"],
            ["Control de acceso", "Completo", "Roles, usuarios, menú modular, hub acceso", "ACC"],
            ["Configuración de nómina", "Completo", "/admin/config — parámetros, festivos", "PAY"],
            ["Configuración de empleados", "Completo", "/admin/employees — CRUD, vínculo usuario", "EMP"],
            ["Tiempo — autoservicio", "Completo", "/my/time — entrada/salida, historial", "TIME"],
            ["Tiempo — administración", "Completo", "/admin/time — CRUD, resolver, corregir", "TADM"],
            ["Reportes y devengados", "Completo", "/reports, /my/reports — Excel, filtros", "RPT"],
            ["Auditoría de tiempo", "Completo", "/admin/time/audit", "AUD"],
            ["Notificaciones admin", "Completo", "Banner incompletos en dashboard", "NOT"],
            ["i18n web (frontend)", "Completo", "es-CO / en-US, selector, formatos", "I18N"],
            ["Operaciones / OSI / móvil", "No implementado", "—", "Fuera de alcance"],
        ],
    )

    # 3. Alcance y exclusiones
    doc.add_heading("3. Alcance y exclusiones", level=1)
    doc.add_heading("3.1 Incluido", level=2)
    for item in [
        "Flujos felices y de error principales por módulo.",
        "Permisos RBAC y aislamiento multi-tenant.",
        "Reglas de negocio visibles: incompletos bloquean reportes, estados OPEN/CLOSED/INCOMPLETE.",
        "Exportación Excel, filtros de periodo, vista capped/uncapped.",
        "i18n es-CO (default) y en-US en UI web.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3.2 Excluido", level=2)
    for item in [
        "Dominio Operations (OSI, bitácora, portal cliente, digest).",
        "Aplicación móvil.",
        "Traducción de mensajes HTTP del backend (incremento futuro).",
        "Pruebas de carga, penetración y compatibilidad exhaustiva de navegadores.",
        "Envío real de correo en notificaciones (job/scheduler puede requerir config adicional).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # 4. Entorno
    doc.add_heading("4. Entorno y prerequisitos", level=1)
    add_table(
        doc,
        ["Componente", "URL / valor", "Notas"],
        [
            ["Frontend", "http://localhost:5173", "npm run dev"],
            ["Security API", "http://localhost:8080", "Auth, users, roles, platform"],
            ["Business API", "http://localhost:8081", "Employees, time, reports, payroll"],
            ["PostgreSQL", "localhost:5432 / stepcore_security", "Docker: stepcore_security_db"],
            ["Seed QA", "scripts/qa-seed-data.sql", "Tenants acme y globex"],
        ],
    )
    doc.add_paragraph(
        "Checklist previo: (1) Docker postgres healthy, (2) backends 8080/8081 arriba, "
        "(3) frontend dev server, (4) seed QA ejecutado, (5) nómina 2026 presente en tenant de prueba."
    )

    # 5. Roles
    doc.add_heading("5. Roles y permisos de prueba", level=1)
    add_table(
        doc,
        ["Rol / perfil", "Tenant ejemplo", "Permisos clave", "Casos típicos"],
        [
            ["EMPLOYEE", "acme", "MY_TIME, MY_PROFILE", "TIME-*, RPT-001, PROF"],
            ["ADMIN", "acme / globex", "Todos ítems payroll + security tenant", "TADM, EMP, PAY, RPT, AUD, NOT"],
            ["PLATFORM_ADMIN", "platform", "PLATFORM_TENANTS", "TEN-001..003"],
            ["Legacy admin", "legacy", "ADMIN", "Regresión tenant histórico"],
        ],
    )

    # 6. Datos de prueba
    doc.add_heading("6. Datos de prueba (QA seed)", level=1)
    doc.add_paragraph(
        "Generados con scripts/qa-seed-data.sql (idempotente). "
        "Contraseña cuentas @*.qa: Admin@2026!"
    )

    doc.add_heading("6.1 Inquilinos", level=2)
    add_table(
        doc,
        ["Empresa", "Slug", "Admin", "Contraseña / notas"],
        [
            ["ACME Corp", "acme", "admin@acme.com", "Provisión original"],
            ["Globex", "globex", "admin@globex.com", "Provisión original"],
            ["Legacy", "legacy", "admin@stepcore.com", "Admin@2026!"],
            ["Plataforma", "platform", "platform@stepcore.com", "Platform@2026!"],
        ],
    )

    doc.add_heading("6.2 ACME — empleados (@acme.qa)", level=2)
    add_table(
        doc,
        ["Nombre", "Correo", "Documento", "Salario COP", "Datos de tiempo sembrados"],
        [
            ["Maria Lopez", "maria.lopez@acme.qa", "CC 1001001001", "2.800.000", "8h normal, overtime, corregido"],
            ["Carlos Ruiz", "carlos.ruiz@acme.qa", "CC 1001001002", "3.200.000", "Nocturno, festivo 25-may"],
            ["Ana Torres", "ana.torres@acme.qa", "CC 1001001003", "2.500.000", "INCOMPLETE 19-may"],
            ["Diego Morales", "diego.morales@acme.qa", "CE 1001001004", "3.500.000", "OPEN hoy, día corto"],
        ],
    )

    doc.add_heading("6.3 Globex — empleados (@globex.qa)", level=2)
    add_table(
        doc,
        ["Nombre", "Correo", "Documento", "Salario COP", "Datos de tiempo sembrados"],
        [
            ["Sofia Ramirez", "sofia.ramirez@globex.qa", "CC 2002002001", "2.900.000", "Normal, extendido, corregido"],
            ["Luis Herrera", "luis.herrera@globex.qa", "CC 2002002002", "3.100.000", "Nocturno"],
            ["Paula Castro", "paula.castro@globex.qa", "TI 2002002003", "2.400.000", "INCOMPLETE 21-may"],
            ["Jorge Mendoza", "jorge.mendoza@globex.qa", "CC 2002002004", "3.600.000", "OPEN hoy, parcial"],
        ],
    )

    doc.add_heading("6.4 Resumen registros de tiempo", level=2)
    add_table(
        doc,
        ["Tenant", "CLOSED", "INCOMPLETE", "OPEN"],
        [["acme", "9", "1", "1"], ["globex", "8", "1", "1"]],
    )
    doc.add_paragraph(
        "Nómina 2026 y festivo 2026-05-25 en ambos tenants. Re-seed: "
        "Get-Content scripts/qa-seed-data.sql | docker exec -i stepcore_security_db "
        "psql -U gmm_user -d stepcore_security"
    )

    # 7. Casos por módulo
    doc.add_heading("7. Casos de prueba por módulo", level=1)

    doc.add_heading("7.1 Autenticación y sesión", level=2)
    add_module_intro(
        doc,
        "Descripción",
        "Valida login multi-tenant con JWT, protección de rutas y gestión básica de sesión.",
        ["/login", "/dashboard", "/my/profile"],
    )
    add_cases(doc, AUTH_CASES, "Autenticación")

    doc.add_heading("7.2 Multi-tenancy y plataforma", level=2)
    add_module_intro(
        doc,
        "Descripción",
        "Administración SaaS de inquilinos y verificación de aislamiento de datos.",
        ["/platform/tenants", "/platform/tenants/new"],
    )
    add_cases(doc, TEN_CASES, "Multi-tenancy")

    doc.add_heading("7.3 Control de acceso", level=2)
    add_module_intro(
        doc,
        "Descripción",
        "Roles, usuarios y permisos de menú modular por tenant.",
        ["/admin/access", "/admin/access/roles", "/admin/access/users", "/admin/access/menu"],
    )
    add_cases(doc, ACC_CASES, "Control de acceso")

    doc.add_heading("7.4 Configuración de nómina", level=2)
    add_module_intro(
        doc,
        "Descripción",
        "Parámetros anuales de devengado colombiano: salarios, rangos horarios, recargos y festivos.",
        ["/admin/config"],
    )
    add_cases(doc, PAY_CASES, "Nómina")

    doc.add_heading("7.5 Configuración de empleados", level=2)
    add_module_intro(
        doc,
        "Descripción",
        "Maestro de empleados separado de usuarios del sistema; requisito para tiempo y reportes.",
        ["/admin/employees", "/admin/employees/new", "/admin/employees/:id"],
    )
    add_cases(doc, EMP_CASES, "Empleados")

    doc.add_heading("7.6 Registro de tiempo — autoservicio", level=2)
    add_module_intro(
        doc,
        "Descripción",
        "Clock-in/out del empleado, historial con devengados capped e indicadores de incompletos.",
        ["/my/time", "/my/reports"],
    )
    add_cases(doc, TIME_CASES, "Tiempo (empleado)")

    doc.add_heading("7.7 Registro de tiempo — administración", level=2)
    add_module_intro(
        doc,
        "Descripción",
        "Gestión admin: listados, incompletos, resolver, corregir, reabrir y crear registros.",
        ["/admin/time"],
    )
    add_cases(doc, TADM_CASES, "Tiempo (admin)")

    doc.add_heading("7.8 Reportes y devengados", level=2)
    add_module_intro(
        doc,
        "Descripción",
        "Reporte A (capped) y B (uncapped), filtros, bloqueo por incompletos, Excel y highlights.",
        ["/reports", "/my/reports"],
    )
    add_cases(doc, RPT_CASES, "Reportes")

    doc.add_heading("7.9 Auditoría de registros de tiempo", level=2)
    add_module_intro(
        doc,
        "Descripción",
        "Trazabilidad de cambios administrativos sobre registros de tiempo.",
        ["/admin/time/audit"],
    )
    add_cases(doc, AUD_CASES, "Auditoría")

    doc.add_heading("7.10 Notificaciones administrativas", level=2)
    add_module_intro(
        doc,
        "Descripción",
        "Banner in-app para admins cuando existen registros incompletos (job + API notifications).",
        ["Dashboard (/dashboard)", "API GET /api/v1/notifications/admin"],
    )
    add_cases(doc, NOT_CASES, "Notificaciones")

    doc.add_heading("7.11 Perfil de usuario", level=2)
    add_cases(doc, PROF_CASES, "Perfil")

    doc.add_heading("7.12 Internacionalización (i18n web)", level=2)
    add_module_intro(
        doc,
        "Descripción",
        "Locale es-CO/en-US, selector, persistencia, menú por código y formatos.",
        ["Todas las pantallas — selector en login y sidebar"],
    )
    add_cases(doc, I18N_CASES, "i18n")

    # 8. Matriz de trazabilidad
    doc.add_heading("8. Matriz de trazabilidad requisitos ↔ casos", level=1)
    add_table(
        doc,
        ["Requisito / regla (requirements-v5)", "Casos relacionados"],
        [
            ["Login JWT + tenant slug", "AUTH-001..003"],
            ["RBAC por permisos de menú", "AUTH-005, ACC-002..003"],
            ["Aislamiento tenant (RLS)", "TEN-005, AUTH-003"],
            ["Parámetros nómina anuales + festivos", "PAY-001..005, RPT-008"],
            ["Empleado vinculado a usuario", "EMP-005, TIME-001"],
            ["Estados OPEN / CLOSED / INCOMPLETE", "TIME-001..002, TADM-002..003"],
            ["Incompletos bloquean reporte (409)", "TIME-005, RPT-004"],
            ["Reporte capped vs uncapped", "RPT-001..003"],
            ["Export Excel", "RPT-007"],
            ["Auditoría admin time records", "AUD-001..004, TADM-004..006"],
            ["Notificación incompletos a admin", "NOT-001..003"],
            ["UI bilingüe es-CO / en-US", "I18N-001..008"],
        ],
    )

    # 9. Matriz pantallas
    doc.add_heading("9. Matriz de cobertura por pantalla", level=1)
    add_table(
        doc,
        ["Ruta", "Módulo", "Casos mínimos recomendados"],
        [
            ["/login", "Auth + i18n", "AUTH-001, I18N-001..002"],
            ["/dashboard", "Notificaciones", "NOT-001, TIME-006"],
            ["/my/time", "Tiempo", "TIME-001..007"],
            ["/my/reports", "Reportes", "RPT-001, RPT-005"],
            ["/reports", "Reportes admin", "RPT-002..007"],
            ["/admin/time", "Tiempo admin", "TADM-001..006"],
            ["/admin/time/audit", "Auditoría", "AUD-001..003"],
            ["/admin/employees", "Empleados", "EMP-001..004"],
            ["/admin/config", "Nómina", "PAY-001..005"],
            ["/admin/access/*", "Acceso", "ACC-001..005"],
            ["/platform/tenants", "Plataforma", "TEN-001..004"],
            ["/my/profile", "Perfil", "PROF-001..002, AUTH-006"],
        ],
    )

    # 10. Criterios de aceptación
    doc.add_heading("10. Criterios de aceptación generales", level=1)
    for item in [
        "≥ 95 % casos de prioridad Alta pasan sin defectos bloqueantes.",
        "Todos los casos TEN-005, AUTH-003, RPT-004 (aislamiento y reglas críticas) pasan.",
        "Export Excel y auditoría operativos en al menos un tenant QA.",
        "i18n: es-CO sin cadenas inglés visibles en flujos principales.",
        "Defectos Medios documentados con workaround o ticket asociado.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # 11. Riesgos
    doc.add_heading("11. Riesgos y limitaciones", level=1)
    add_table(
        doc,
        ["Item", "Impacto", "Notas"],
        [
            ["Registros OPEN atados a fecha 'hoy' en seed", "Medio", "Re-ejecutar seed periódicamente"],
            ["Mensajes API en inglés", "Medio", "Incremento feature-i18n-api-messages pendiente"],
            ["Notificaciones: título/cuerpo desde BD", "Bajo", "No traducidos en frontend"],
            ["Job nocturno incompletos", "Medio", "Verificar scheduler si NOT-001 no aparece"],
            ["Contraseña admin acme/globex desconocida", "Alto", "Reset vía platform o SQL"],
        ],
    )

    # 12. Registro ejecución
    doc.add_heading("12. Registro de ejecución", level=1)
    add_table(
        doc,
        ["ID", "Ejecutor", "Fecha", "Build", "Resultado", "Defecto / observación"],
        [
            ["AUTH-001", "", "", "", "Pendiente", ""],
            ["TIME-001", "", "", "", "Pendiente", ""],
            ["RPT-007", "", "", "", "Pendiente", ""],
            ["I18N-001", "", "", "", "Pendiente", ""],
        ],
    )
    doc.add_paragraph("Ampliar filas para cada caso ejecutado en el ciclo QA.")

    # 13. Referencias
    doc.add_heading("13. Referencias", level=1)
    for ref in [
        "docs/architecture.md — mapa de módulos",
        "requirements/requirements-v5.md — reglas de negocio",
        "scripts/QA-SEED-CREDENTIALS.md",
        "scripts/qa-seed-data.sql",
        "openspec/specs/time-record-audit-trail/spec.md",
        "openspec/specs/multi-tenancy/spec.md",
        "PR Frontend i18n: github.com/jfranciscogomezn/gmm-devengos-frontend/pull/11",
    ]:
        doc.add_paragraph(ref, style="List Bullet")

    # Resumen conteo
    total = (
        len(AUTH_CASES) + len(TEN_CASES) + len(ACC_CASES) + len(PAY_CASES)
        + len(EMP_CASES) + len(TIME_CASES) + len(TADM_CASES) + len(RPT_CASES)
        + len(AUD_CASES) + len(NOT_CASES) + len(PROF_CASES) + len(I18N_CASES)
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Total casos de prueba documentados: {total}").bold = True

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT} ({len(AUTH_CASES)+len(TEN_CASES)+len(ACC_CASES)+len(PAY_CASES)+len(EMP_CASES)+len(TIME_CASES)+len(TADM_CASES)+len(RPT_CASES)+len(AUD_CASES)+len(NOT_CASES)+len(PROF_CASES)+len(I18N_CASES)} test cases)")


if __name__ == "__main__":
    main()
